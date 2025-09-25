"""
Document processor for knowledge base ingestion
"""

import os
import json
import hashlib
import tiktoken
from pathlib import Path
from typing import List, Dict, Any, Optional
import openai
from openai import OpenAI
from datetime import datetime
from sqlalchemy.orm import Session
from backend.db.database import get_db
from backend.db.models import Document, DocumentChunk
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentProcessor:
    """Processes documents for knowledge base ingestion."""

    def __init__(self, documents_dir: str = "documents"):
        self.documents_dir = Path(documents_dir)

        # Load environment variables
        load_dotenv(dotenv_path=Path(__file__).parent.parent.parent / ".env")

        # Initialize OpenAI client
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            self.openai_client = OpenAI(api_key=api_key)
        else:
            self.openai_client = None
            print("Warning: OPENAI_API_KEY is not set. Document embedding will not work.")

        # Initialize tokenizer for chunking
        self.tokenizer = tiktoken.get_encoding("cl100k_base")

        # Chunking parameters
        self.chunk_size = 1000  # tokens
        self.chunk_overlap = 200  # tokens
    
    def process_document(self, file_path: Path) -> Dict[str, Any]:
        """
        Process a single document and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing processed document data
        """
        try:
            # Read file content based on file type
            content = self._read_file_content(file_path)
            
            # Clean and preprocess content
            cleaned_content = self._clean_content(content)
            
            # Chunk the content
            chunks = self._chunk_text(cleaned_content)
            
            # Generate metadata
            metadata = {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_type": file_path.suffix,
                "file_size": file_path.stat().st_size,
                "processed_at": datetime.utcnow().isoformat(),
                "chunk_count": len(chunks)
            }
            
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_type": file_path.suffix,
                "content": cleaned_content,
                "chunks": chunks,
                "metadata": metadata,
                "processed_at": metadata["processed_at"]
            }
            
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            return None
    
    def _read_file_content(self, file_path: Path) -> str:
        """Read content from various file types."""
        file_extension = file_path.suffix.lower()
        
        if file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.yaml', '.yml']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_extension in ['.pdf']:
            if PDF_AVAILABLE:
                try:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except Exception as e:
                    return f"Error reading PDF {file_path.name}: {str(e)}"
            else:
                return f"PDF content from {file_path.name} - PyPDF2 not installed"
        elif file_extension in ['.docx']:
            if DOCX_AVAILABLE:
                try:
                    doc = Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except Exception as e:
                    return f"Error reading DOCX {file_path.name}: {str(e)}"
            else:
                return f"DOCX content from {file_path.name} - python-docx not installed"
        else:
            # Try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except:
                return f"Binary file {file_path.name} - content not readable"
    
    def _clean_content(self, content: str) -> str:
        """Clean and preprocess text content."""
        # Remove excessive whitespace
        content = ' '.join(content.split())
        
        # Remove special characters that might interfere with processing
        content = content.replace('\x00', '')  # Remove null bytes
        
        return content.strip()
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks for optimal embedding size."""
        # Tokenize the text
        tokens = self.tokenizer.encode(text)
        
        chunks = []
        start = 0
        
        while start < len(tokens):
            # Calculate end position
            end = min(start + self.chunk_size, len(tokens))
            
            # Extract chunk tokens
            chunk_tokens = tokens[start:end]
            
            # Decode back to text
            chunk_text = self.tokenizer.decode(chunk_tokens)
            chunks.append(chunk_text)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            
            # Prevent infinite loop
            if start >= len(tokens):
                break
        
        return chunks
    
    def generate_embeddings(self, content: str) -> List[float]:
        """
        Generate embeddings for the given content using OpenAI.
        
        Args:
            content: Text content to embed
            
        Returns:
            List of embedding vectors
        """
        try:
            response = self.openai_client.embeddings.create(
                model="text-embedding-3-small",
                input=content
            )
            return response.data[0].embedding
        except Exception as e:
            print(f"Error generating embeddings: {str(e)}")
            return None
    
    def upsert_to_vector_db(self, document_data: Dict[str, Any], db: Session):
        """
        Upsert document chunks to PostgreSQL vector database.

        Args:
            document_data: Processed document data
            db: Database session
        """
        if not document_data or not document_data.get("chunks"):
            return

        try:
            # Check if document already exists
            existing_doc = db.query(Document).filter(Document.file_path == document_data["file_path"]).first()
            if existing_doc:
                print(f"Document {document_data['file_name']} already exists, skipping")
                return

            # Create document record
            doc = Document(
                file_path=document_data["file_path"],
                file_name=document_data["file_name"],
                file_type=document_data["file_type"],
                file_size=document_data["metadata"]["file_size"],
                processed_at=document_data["processed_at"],
                content=document_data["content"]
            )
            db.add(doc)
            db.flush()  # Get the document ID

            # Generate embeddings for each chunk
            for i, chunk in enumerate(document_data["chunks"]):
                # Generate embedding
                embedding = self.generate_embeddings(chunk)
                if embedding is None:
                    continue

                # Prepare metadata
                chunk_metadata = {
                    "file_name": document_data["file_name"],
                    "file_path": document_data["file_path"],
                    "file_type": document_data["file_type"],
                    "chunk_index": i,
                    "chunk_size": len(chunk),
                    "processed_at": document_data["processed_at"]
                }

                # Create chunk record
                chunk_record = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=i,
                    content=chunk,
                    embedding=embedding,
                    chunk_metadata=chunk_metadata
                )
                db.add(chunk_record)

            db.commit()
            print(f"Upserted {len(document_data['chunks'])} chunks from {document_data['file_name']}")

        except Exception as e:
            db.rollback()
            print(f"Error upserting to vector DB: {str(e)}")
    
    def search_similar_documents(self, query: str, n_results: int = 5, db: Session = None) -> List[Dict[str, Any]]:
        """
        Search for similar documents in the PostgreSQL vector database.

        Args:
            query: Search query
            n_results: Number of results to return
            db: Database session

        Returns:
            List of similar documents with metadata and scores
        """
        if db is None:
            # Get a database session
            from backend.db.database import SessionLocal
            db = SessionLocal()

        try:
            # Generate embedding for query
            query_embedding = self.generate_embeddings(query)
            if query_embedding is None:
                return []

            # Search using pgvector cosine similarity
            from sqlalchemy import text
            results = db.query(
                DocumentChunk,
                DocumentChunk.embedding.cosine_distance(query_embedding).label('distance')
            ).join(Document).order_by(
                DocumentChunk.embedding.cosine_distance(query_embedding)
            ).limit(n_results).all()

            # Format results
            similar_docs = []
            for chunk, distance in results:
                doc = {
                    "id": chunk.id,
                    "content": chunk.content,
                    "metadata": chunk.chunk_metadata,
                    "relevance_score": 1 - distance  # Convert distance to similarity
                }
                similar_docs.append(doc)

            return similar_docs

        except Exception as e:
            print(f"Error searching documents: {str(e)}")
            return []
        finally:
            if db:
                db.close()
    
    def process_all_documents(self, db: Session = None):
        """Process all documents in the documents directory and subdirectories."""
        if db is None:
            from backend.db.database import SessionLocal
            db = SessionLocal()

        try:
            if not self.documents_dir.exists():
                print(f"Documents directory {self.documents_dir} does not exist")
                return

            processed_count = 0
            # Recursively find all supported files
            supported_extensions = ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.yaml', '.yml']

            # Add PDF and DOCX if libraries are available
            if PDF_AVAILABLE:
                supported_extensions.append('.pdf')
            if DOCX_AVAILABLE:
                supported_extensions.append('.docx')

            for file_path in self.documents_dir.rglob('*'):
                if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                    print(f"Processing: {file_path.relative_to(self.documents_dir)}")
                    document_data = self.process_document(file_path)
                    if document_data:
                        self.upsert_to_vector_db(document_data, db)
                        processed_count += 1
                        print(f"Processed: {file_path.relative_to(self.documents_dir)}")

            print(f"Total documents processed: {processed_count}")
        finally:
            if db:
                db.close()

if __name__ == "__main__":
    processor = DocumentProcessor()
    processor.process_all_documents()

